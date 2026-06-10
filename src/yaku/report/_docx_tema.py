#!/usr/bin/env python3
"""docx_helpers — construir documentos Word con tema azul/agua + export a PDF.

Helper reutilizable para informes/tutoriales/guías. Requiere `python-docx`.
Conversión a PDF con LibreOffice (macOS). Diseño limpio, SIN emojis (LibreOffice
los rompe en el PDF).

Uso minimo:
    from docx_helpers import DocBuilder, docx_to_pdf
    d = DocBuilder()
    d.cover("Titulo", "Subtitulo", "Autor - Empresa - Fecha")
    d.h1("Introduccion"); d.p("Texto.")
    d.info_box("Nota", "Algo importante.")
    d.table(["Col A", "Col B"], [["1", "2"], ["3", "4"]])
    d.image("figura.png", "Figura 1. Descripcion.")
    d.save("salida.docx"); docx_to_pdf("salida.docx")
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- Paleta azul/agua (cambia aqui para otro tema) ---
AZUL_PROFUNDO = RGBColor(0x0D, 0x47, 0xA1)
AZUL = RGBColor(0x15, 0x65, 0xC0)
AZUL_MEDIO = RGBColor(0x1E, 0x88, 0xE5)
GRIS = RGBColor(0x37, 0x47, 0x4F)
GRIS_SUAVE = RGBColor(0x60, 0x70, 0x78)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
# hex sin '#'
BANNER = "1565C0"
BANNER_ACENTO = "64B5F6"
AGUA_FONDO = "E8F2FC"
ZEBRA = "F4F8FD"
HEADER_TABLA = "1E6FB8"
CODE_BG = "ECEFF1"
LINEA_AZUL = "1565C0"

_SOFFICE = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "soffice",
    "libreoffice",
]


# --- utilidades XML ---
def _shade(cell, hex_fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def _bottom_border(paragraph, color: str = LINEA_AZUL, size: int = 8) -> None:
    """Regla horizontal fina bajo un párrafo (encabezados)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def _left_border(cell, color: str = LINEA_AZUL, size: int = 24) -> None:
    """Borde izquierdo grueso en una celda (acento de caja de nota)."""
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)
    borders.append(left)
    tcpr.append(borders)


def _no_cell_borders(table) -> None:
    """Quita los bordes de una tabla (para banners/cajas de una celda)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


class DocBuilder:
    """Constructor de documentos Word con tema azul/agua (diseño limpio)."""

    def __init__(self, font: str = "Calibri", size: int = 11):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = font
        normal.font.size = Pt(size)
        self._h1 = 0
        self._num = 0           # contador de lista numerada (se reinicia por grupo)
        self._last_numbered = False

    # --- Portada ---
    def cover(self, title: str, subtitle: str = "", footer: str = "") -> "DocBuilder":
        banner = self.doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_cell_borders(banner)
        cell = banner.cell(0, 0)
        _shade(cell, BANNER)
        cell.text = ""
        cell.add_paragraph()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(32)
        r.font.color.rgb = BLANCO
        if subtitle:
            ps = cell.add_paragraph()
            ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = ps.add_run(subtitle)
            rs.font.size = Pt(14)
            rs.font.color.rgb = RGBColor(0xE3, 0xF2, 0xFD)
        cell.add_paragraph()
        # línea de acento bajo el banner
        acc = self.doc.add_paragraph()
        _bottom_border(acc, BANNER_ACENTO, size=18)
        if footer:
            pf = self.doc.add_paragraph()
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rf = pf.add_run(footer)
            rf.font.size = Pt(11)
            rf.font.color.rgb = GRIS_SUAVE
            rf.italic = True
        self.doc.add_paragraph()
        return self

    # --- Encabezados (párrafos estilizados, sin doble estilo de Word) ---
    def h1(self, text: str, numbered: bool = True) -> "DocBuilder":
        self._last_numbered = False
        if numbered:
            self._h1 += 1
            text = f"{self._h1}.  {text}"
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(17)
        r.font.color.rgb = AZUL
        _bottom_border(p, LINEA_AZUL, size=6)
        return self

    def h2(self, text: str) -> "DocBuilder":
        self._last_numbered = False
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = AZUL_MEDIO
        return self

    # --- Texto ---
    def p(self, text: str, size: int = 10.5) -> "DocBuilder":
        self._last_numbered = False
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.08
        r = para.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = GRIS
        return self

    def bullet(self, text: str) -> "DocBuilder":
        self._last_numbered = False
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRIS
        return self

    def numbered(self, text: str) -> "DocBuilder":
        # Contador propio (reinicia al empezar una lista nueva) para evitar que Word
        # continue la numeracion entre secciones (20, 21, 22...).
        if not self._last_numbered:
            self._num = 0
        self._num += 1
        self._last_numbered = True
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        rn = p.add_run(f"{self._num}.  ")
        rn.bold = True
        rn.font.size = Pt(10.5)
        rn.font.color.rgb = AZUL_MEDIO
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRIS
        return self

    # --- Caja de nota (borde azul a la izquierda, fondo agua) ---
    def info_box(self, title: str, text: str) -> "DocBuilder":
        self._last_numbered = False
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_cell_borders(table)
        body = table.cell(0, 0)
        _shade(body, AGUA_FONDO)
        _left_border(body, BANNER, size=30)
        body.text = ""
        p = body.paragraphs[0]
        rt = p.add_run(title)
        rt.bold = True
        rt.font.color.rgb = AZUL_PROFUNDO
        rt.font.size = Pt(10.5)
        p2 = body.add_paragraph()
        r2 = p2.add_run(text)
        r2.font.color.rgb = GRIS
        r2.font.size = Pt(10)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return self

    # --- Tabla con cabecera azul y filas zebra ---
    def table(self, header: list[str], rows: list[list[str]]) -> "DocBuilder":
        self._last_numbered = False
        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, htext in enumerate(header):
            c = t.rows[0].cells[j]
            _shade(c, HEADER_TABLA)
            c.text = ""
            run = c.paragraphs[0].add_run(htext)
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(10)
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            fill = ZEBRA if i % 2 == 0 else "FFFFFF"
            for j, val in enumerate(row):
                _shade(cells[j], fill)
                cells[j].text = ""
                run = cells[j].paragraphs[0].add_run(str(val))
                run.font.size = Pt(9.5)
                run.bold = j == 0
                run.font.color.rgb = AZUL_PROFUNDO if j == 0 else GRIS
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return self

    # --- Bloque de codigo ---
    def code(self, code: str) -> "DocBuilder":
        self._last_numbered = False
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0)
        _shade(cell, CODE_BG)
        cell.text = ""
        for i, line in enumerate(code.strip("\n").split("\n")):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line or " ")
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = AZUL_PROFUNDO
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return self

    # --- Imagen con pie ---
    def image(self, path, caption: str = "", width_in: float = 5.2) -> "DocBuilder":
        self._last_numbered = False
        path = Path(path)
        if not path.exists():
            return self
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.add_run().add_picture(str(path), width=Inches(width_in))
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = GRIS_SUAVE
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return self

    def page_break(self) -> "DocBuilder":
        self._last_numbered = False
        self.doc.add_page_break()
        return self

    def save(self, path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path


def docx_to_pdf(docx_path, out_dir=None) -> Path:
    """Convierte un .docx a PDF con LibreOffice (alta fidelidad). Devuelve la ruta del PDF."""
    docx_path = Path(docx_path)
    out_dir = Path(out_dir) if out_dir else docx_path.parent
    last_err = None
    for soffice in _SOFFICE:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
                check=True, capture_output=True, text=True,
            )
            pdf = out_dir / (docx_path.stem + ".pdf")
            if pdf.exists():
                return pdf
        except Exception as exc:
            last_err = exc
    raise RuntimeError(f"No se pudo convertir a PDF (LibreOffice). Ultimo error: {last_err}")
